from docx import Document
Document("book.docx")
doc=Document("book.docx")
count=0
for p in doc.paragraphs:
    if p.text.strip():
        count=count+1
print("Total non-empty lines:", count) 