from docx import Document

doc = Document("book.docx")   # make sure the filename matches
paragraphs = [p.text for p in doc.paragraphs]

# Remove empty lines
lines = [line.strip() for line in paragraphs if line.strip()]

print("Total non-empty lines:", len(lines))
for p in doc.paragraphs:
    if p.text.strip():
        print("First non-empty paragraph:", p.text)
        break